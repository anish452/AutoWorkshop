"""
Generate Master's Final Report for AutoRepairAgent (16,000–20,000 words).
Output: docs/Final_Report_AutoRepairAgent.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE = Path(__file__).resolve().parent
OUTPUT = BASE / "Final_Report_AutoRepairAgent.docx"
SCREENSHOTS_R1 = BASE / "presentation-screenshots"
SCREENSHOTS_R2 = BASE / "presentation-screenshots-report2"


def setup_styles(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 16), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(15, 39, 68)


def add_title_page(doc):
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AutoRepairAgent")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(15, 39, 68)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "An AI-Powered Web-Based Management System for Vehicle Repair Workshops"
    )
    sr.font.size = Pt(14)
    sr.italic = True

    doc.add_paragraph()
    fin = doc.add_paragraph()
    fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fin.add_run("Final Report\n").bold = True
    fin.add_run("Specialised Project — Master's Level\n\n")
    fin.add_run("[Student Full Name]\n")
    fin.add_run("[Student ID Number]\n\n")
    fin.add_run("[Course Name / Programme]\n")
    fin.add_run("[College / University Name]\n\n")
    fin.add_run("Supervisor: [Supervisor Name]\n\n")
    fin.add_run("Submission Date: July 2026\n")
    doc.add_page_break()


def add_abstract(doc):
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Vehicle repair workshops face persistent operational challenges including manual "
        "complaint triage, inconsistent job routing across specialist departments, limited "
        "visibility of work-in-progress for customers and managers, and fragmented record-keeping "
        "across customers, vehicles, and repair tasks. This project designed, implemented, and "
        "evaluated AutoRepairAgent, a full-stack web application that integrates artificial "
        "intelligence into the complaint intake and job creation workflow for automotive repair centres."
    )
    doc.add_paragraph(
        "The system comprises a Node.js and Express.js REST API backed by PostgreSQL and Prisma ORM, "
        "a React single-page application built with Material UI, and a DeepSeek large language model "
        "integration for multilingual complaint classification. Clean Architecture and the Repository "
        "pattern were adopted to separate presentation, application, infrastructure, and domain concerns. "
        "Role-based access control supports eight distinct user roles spanning administration, job advisory, "
        "five workshop departments, and customer portal access."
    )
    doc.add_paragraph(
        "Development proceeded in two phases. Phase 1 established the database schema, authentication, "
        "administrative modules, and front-end design framework. Phase 2 delivered customer and vehicle "
        "management, AI-powered job creation, automatic department assignment, job lifecycle management, "
        "AI analysis logging, and role-specific dashboards. Functional testing using seeded data and manual "
        "end-to-end scenarios confirmed that a single complaint can be decomposed into multiple department-specific "
        "jobs with confidence metadata and audit trails."
    )
    doc.add_paragraph(
        "Results demonstrate that the project objectives were substantially achieved within the academic timeframe. "
        "Limitations include the absence of production deployment, automated test coverage, multilingual user "
        "interface localisation, and naive first-available technician assignment. Future work should prioritise "
        "cloud deployment, Sinhala and Tamil interface support, notification services, automated testing pipelines, "
        "and load-balanced assignment algorithms."
    )
    doc.add_paragraph()
    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run(
        "vehicle repair management; artificial intelligence; complaint classification; "
        "role-based access control; full-stack web application; DeepSeek; Clean Architecture"
    )
    doc.add_page_break()


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def word_count_estimate(doc):
    total = 0
    for p in doc.paragraphs:
        total += len(p.text.split())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += len(cell.text.split())
    return total


def build_report(doc):
    # ── TABLE OF CONTENTS placeholder ───────────────────────────────────────
    add_heading(doc, "Table of Contents")
    toc_items = [
        "Abstract ........................................................... ii",
        "List of Figures .................................................... iv",
        "List of Tables ..................................................... v",
        "1. Introduction .................................................... 1",
        "   1.1 Background and Context",
        "   1.2 Problem Statement",
        "   1.3 Motivation and Significance",
        "   1.4 Research Questions",
        "   1.5 Report Structure",
        "2. Literature Review ................................................ 8",
        "   2.1 Vehicle Repair Workshop Management Systems",
        "   2.2 Artificial Intelligence in Automotive Service",
        "   2.3 Web Application Architectures and Security",
        "   2.4 Role-Based Access Control in Enterprise Systems",
        "   2.5 Synthesis and Research Gap",
        "3. Project Objectives and Scope .................................... 22",
        "   3.1 Project Objectives",
        "   3.2 Success Criteria",
        "   3.3 Project Scope",
        "   3.4 Out of Scope",
        "   3.5 Assumptions and Constraints",
        "4. Methodology ..................................................... 30",
        "   4.1 Development Approach",
        "   4.2 Requirements Gathering",
        "   4.3 System Design Methodology",
        "   4.4 Implementation Methodology",
        "   4.5 Testing and Validation Approach",
        "   4.6 Ethical Considerations",
        "5. System Design and Implementation ............................... 42",
        "   5.1 Architecture Overview",
        "   5.2 Database Design",
        "   5.3 Authentication and Security",
        "   5.4 Backend Modules",
        "   5.5 AI Integration",
        "   5.6 Frontend Design",
        "   5.7 Development Environment",
        "6. Results and Outcomes ........................................... 62",
        "   6.1 Phase 1 Outcomes",
        "   6.2 Phase 2 Outcomes",
        "   6.3 Functional Workflow Results",
        "   6.4 User Interface Outcomes",
        "   6.5 Security and Audit Outcomes",
        "   6.6 Quantitative Summary",
        "7. Discussion ...................................................... 74",
        "   7.1 Alignment with Objectives",
        "   7.2 Critical Evaluation of AI Classification",
        "   7.3 Architectural Decisions",
        "   7.4 Limitations and Shortcomings",
        "   7.5 Threats to Validity",
        "8. Conclusions and Future Work .................................... 86",
        "   8.1 Conclusions",
        "   8.2 Contributions",
        "   8.3 Future Work",
        "References ........................................................ 92",
        "Appendices ........................................................ 96",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    add_heading(doc, "List of Figures")
    figures = [
        "Figure 5.1  Three-tier system architecture with Clean Architecture layers",
        "Figure 5.2  Entity-relationship model of the PostgreSQL database schema",
        "Figure 5.3  Job lifecycle state machine",
        "Figure 5.4  AI complaint analysis sequence diagram",
        "Figure 6.1  Login page — authentication interface",
        "Figure 6.2  Admin users management page",
        "Figure 6.3  Department management page",
        "Figure 6.4  Customer management page",
        "Figure 6.5  Vehicle registration page",
        "Figure 6.6  AI job creation page",
        "Figure 6.7  Jobs list with department routing",
        "Figure 6.8  AI analysis log viewer",
        "Figure 6.9  Admin dashboard with KPI cards",
    ]
    for f in figures:
        doc.add_paragraph(f)
    doc.add_page_break()

    add_heading(doc, "List of Tables")
    tables = [
        "Table 3.1  Project objectives and completion status",
        "Table 3.2  In-scope and out-of-scope features",
        "Table 4.1  Technology stack",
        "Table 5.1  Database tables and descriptions",
        "Table 5.2  API endpoint summary (34 endpoints)",
        "Table 5.3  Role-based access control matrix",
        "Table 6.1  Phase 1 deliverables",
        "Table 6.2  Phase 2 deliverables",
        "Table 6.3  Test accounts and seeded data",
        "Table 7.1  Limitations and impact assessment",
    ]
    for t in tables:
        doc.add_paragraph(t)
    doc.add_page_break()

    # Import chapter content
    from final_report_chapters import (
        chapter1_introduction,
        chapter2_literature,
        chapter3_objectives,
        chapter4_methodology,
        chapter5_implementation,
        chapter6_results,
        chapter7_discussion,
        chapter8_conclusions,
        references,
        appendices,
    )

    chapter1_introduction(doc)
    chapter2_literature(doc)
    chapter3_objectives(doc)
    chapter4_methodology(doc)
    chapter5_implementation(doc)
    chapter6_results(doc)
    chapter7_discussion(doc)
    chapter8_conclusions(doc)
    references(doc)
    appendices(doc)


def main():
    doc = Document()
    setup_styles(doc)
    sections = doc.sections[0]
    sections.top_margin = Inches(1)
    sections.bottom_margin = Inches(1)
    sections.left_margin = Inches(1.25)
    sections.right_margin = Inches(1)

    add_title_page(doc)
    add_abstract(doc)
    build_report(doc)

    doc.save(OUTPUT)
    wc = word_count_estimate(doc)
    print(f"Created: {OUTPUT}")
    print(f"Estimated word count: {wc:,}")
    if wc < 16000:
        print("WARNING: Below 16,000 word target — expand chapters as needed.")
    elif wc > 20000:
        print("NOTE: Above 20,000 word target.")


if __name__ == "__main__":
    main()
